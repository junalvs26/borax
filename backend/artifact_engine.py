import os
import re
import json
import csv
import uuid
from typing import Dict, Any, List, Optional

# --- DOCX Imports ---
import docx
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

# --- XLSX Imports ---
try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

# --- PDF Imports (ReportLab) ---
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether, HRFlowable
    from reportlab.pdfgen import canvas
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False


class NumberedCanvas(canvas.Canvas):
    """ReportLab Canvas extension for dynamic 'Página X de Y' footer and running header."""
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_decorations(num_pages)
            super().showPage()
        super().save()

    def draw_page_decorations(self, page_count):
        if self._pageNumber == 1:
            # Skip header and footer on cover page
            return

        self.saveState()
        self.setFont("Helvetica", 9)
        self.setFillColor(colors.HexColor("#71717A"))

        # Running Header
        self.drawString(54, 800, "PLATAFORMA BORAX - DOCUMENTO DE ALTA PERFORMANCE")
        self.setStrokeColor(colors.HexColor("#27272A"))
        self.setLineWidth(0.5)
        self.line(54, 792, 541, 792)

        # Running Footer
        page_str = f"Página {self._pageNumber} de {page_count}"
        self.drawRightString(541, 36, page_str)
        self.drawString(54, 36, "Confidencial - Gerado via Engine Local BORAX")
        self.line(54, 48, 541, 48)

        self.restoreState()


class ArtifactEngine:
    def __init__(self):
        pass

    def compile_docx(self, markdown_text: str, filepath: Optional[str] = None, title: str = "Trabalho Acadêmico") -> str:
        """
        Compiles long Markdown into a fully compliant ABNT Word (.docx) document.
        - Margins: Top 3cm, Left 3cm, Bottom 2cm, Right 2cm.
        - Font: Arial 12pt, 1.5 line spacing, 1.25cm first line indent, Justified.
        - Native Word Tables with dark header styling and clean borders.
        """
        if not filepath:
            exports_dir = os.path.expanduser("~/.borax/exports")
            os.makedirs(exports_dir, exist_ok=True)
            safe_title = re.sub(r'[^\w\s-]', '', title).strip().replace(" ", "_").lower() or "documento_abnt"
            filepath = os.path.join(exports_dir, f"{safe_title}_{uuid.uuid4().hex[:6]}.docx")

        doc = docx.Document()

        # Set ABNT Page Margins
        for section in doc.sections:
            section.top_margin = Cm(3.0)
            section.left_margin = Cm(3.0)
            section.bottom_margin = Cm(2.0)
            section.right_margin = Cm(2.0)

        # Set Normal Paragraph Style (ABNT Body Text)
        style_normal = doc.styles['Normal']
        font = style_normal.font
        font.name = 'Arial'
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)
        style_normal.paragraph_format.line_spacing = 1.5
        style_normal.paragraph_format.space_after = Pt(6)
        style_normal.paragraph_format.first_line_indent = Cm(1.25)
        style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        lines = markdown_text.split("\n")
        in_table = False
        table_buffer = []

        def flush_table_buffer(buffer):
            if not buffer:
                return
            rows_data = []
            for t_line in buffer:
                if t_line.strip().startswith("|") and "---" not in t_line:
                    cells = [c.strip() for c in t_line.strip().split("|")[1:-1]]
                    if cells:
                        rows_data.append(cells)
            
            if not rows_data:
                return

            col_count = max(len(r) for r in rows_data)
            table = doc.add_table(rows=len(rows_data), cols=col_count)
            table.autofit = False

            for r_idx, row in enumerate(rows_data):
                is_header = (r_idx == 0)
                for c_idx, cell_value in enumerate(row):
                    if c_idx < col_count:
                        cell = table.cell(r_idx, c_idx)
                        cell.text = cell_value
                        
                        # Style paragraph inside cell
                        p = cell.paragraphs[0]
                        p.paragraph_format.first_line_indent = Cm(0)
                        p.paragraph_format.line_spacing = 1.15
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        
                        for run in p.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(10)
                            if is_header:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)

                        # Cell background shading
                        if is_header:
                            shd = parse_xml(r'<w:shd {} w:fill="14111F"/>'.format(nsdecls('w')))
                            cell._tc.get_or_add_tcPr().append(shd)
                        elif r_idx % 2 == 1:
                            shd = parse_xml(r'<w:shd {} w:fill="F4F4F5"/>'.format(nsdecls('w')))
                            cell._tc.get_or_add_tcPr().append(shd)

            # Table space after
            p_space = doc.add_paragraph()
            p_space.paragraph_format.first_line_indent = Cm(0)
            p_space.paragraph_format.space_after = Pt(6)

        for line in lines:
            stripped = line.strip()

            # Handle Markdown Tables
            if stripped.startswith("|"):
                in_table = True
                table_buffer.append(stripped)
                continue
            else:
                if in_table:
                    flush_table_buffer(table_buffer)
                    table_buffer = []
                    in_table = False

            if not stripped:
                continue

            # Headings
            if stripped.startswith("# "):
                h_text = stripped[2:].strip()
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(h_text.upper())
                run.font.name = 'Arial'
                run.font.size = Pt(14)
                run.font.bold = True
                run.font.color.rgb = RGBColor(20, 17, 31)

            elif stripped.startswith("## "):
                h_text = stripped[3:].strip()
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(h_text)
                run.font.name = 'Arial'
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

            elif stripped.startswith("### "):
                h_text = stripped[4:].strip()
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(h_text)
                run.font.name = 'Arial'
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.italic = True

            elif stripped.startswith("- ") or stripped.startswith("* "):
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                text = stripped[2:].strip()
                self._append_formatted_runs(p, text)
            else:
                p = doc.add_paragraph()
                self._append_formatted_runs(p, stripped)

        if in_table:
            flush_table_buffer(table_buffer)

        doc.save(filepath)
        return filepath

    def _append_formatted_runs(self, paragraph, text: str):
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                r = paragraph.add_run(part[2:-2])
                r.font.name = 'Arial'
                r.font.size = Pt(12)
                r.bold = True
            else:
                r = paragraph.add_run(part)
                r.font.name = 'Arial'
                r.font.size = Pt(12)

    # =========================================================================
    # 2. STRATEGY B: PLANILHAS (XLSX via openpyxl com Fórmulas Reais)
    # =========================================================================
    def compile_xlsx(self, content_str: str, filepath: str, title: str = "Planilha de Dados") -> str:
        """
        Generates a professional Excel (.xlsx) file using openpyxl with active formulas.
        Parses JSON array/object or CSV/Markdown Table from LLM output.
        - Dark purple header styling (#14111F), bold white text.
        - Preserves and evaluates Excel formulas (e.g. '=SUM(B2:B10)').
        - Auto-adjusts column widths.
        - Number formatting (currency, percentage, numbers).
        """
        if not OPENPYXL_AVAILABLE:
            raise ImportError("openpyxl não está instalado no ambiente backend Python.")

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = title[:30]
        ws.views.sheetView[0].showGridLines = True

        headers = []
        rows = []

        # Try parsing JSON structure first
        try:
            clean_json = content_str.strip()
            if clean_json.startswith("```json"):
                clean_json = clean_json[7:]
            if clean_json.endswith("```"):
                clean_json = clean_json[:-3]
            clean_json = clean_json.strip()

            parsed_data = json.loads(clean_json)
            if isinstance(parsed_data, dict):
                headers = parsed_data.get("headers", [])
                rows = parsed_data.get("rows", [])
            elif isinstance(parsed_data, list) and len(parsed_data) > 0:
                headers = list(parsed_data[0].keys())
                rows = [[item.get(h, "") for h in headers] for item in parsed_data]
        except Exception:
            pass

        # Fallback: Parse CSV / Markdown Table
        if not headers or not rows:
            lines = [l.strip() for l in content_str.split("\n") if l.strip()]
            table_lines = [l for l in lines if l.startswith("|") and "---" not in l]
            if table_lines:
                headers = [c.strip() for c in table_lines[0].split("|")[1:-1]]
                for t_line in table_lines[1:]:
                    r_cells = [c.strip() for c in t_line.split("|")[1:-1]]
                    if r_cells:
                        rows.append(r_cells)

        # Final Default Fallback if parsing fails
        if not headers:
            headers = ["Item / Categoria", "Quantidade", "Valor Unitário (R$)", "Total (R$)"]
            rows = [
                ["Desenvolvimento de Software", 120, 150.00, "=B2*C2"],
                ["Consultoria e Arquitetura", 40, 200.00, "=B3*C3"],
                ["Infraestrutura Local", 1, 5000.00, "=B4*C4"],
                ["Total Geral", "", "", "=SUM(D2:D4)"]
            ]

        # Styles
        header_font = Font(name="Arial", size=11, bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="14111F", end_color="14111F", fill_type="solid")
        data_font = Font(name="Arial", size=10)
        bold_font = Font(name="Arial", size=10, bold=True)

        thin_border = Border(
            left=Side(style='thin', color='E4E4E7'),
            right=Side(style='thin', color='E4E4E7'),
            top=Side(style='thin', color='E4E4E7'),
            bottom=Side(style='thin', color='E4E4E7')
        )
        zebra_fill = PatternFill(start_color="F8F8FA", end_color="F8F8FA", fill_type="solid")

        # Write Headers
        ws.append(headers)
        for col_idx in range(1, len(headers) + 1):
            cell = ws.cell(row=1, column=col_idx)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = thin_border

        ws.row_dimensions[1].height = 26

        # Write Data Rows
        for r_idx, row_data in enumerate(rows, start=2):
            ws.append(row_data)
            ws.row_dimensions[r_idx].height = 20
            is_zebra = (r_idx % 2 == 1)

            for col_idx, cell_value in enumerate(row_data, start=1):
                cell = ws.cell(row=r_idx, column=col_idx)
                cell.font = bold_font if str(cell_value).startswith("=") or "total" in str(row_data[0]).lower() else data_font
                cell.border = thin_border

                if is_zebra:
                    cell.fill = zebra_fill

                # Format Numbers / Currency / Formulas
                val_str = str(cell_value).strip()
                if val_str.startswith("="):
                    cell.value = val_str
                elif re.match(r'^-?\d+(\.\d+)?$', val_str):
                    num_val = float(val_str)
                    cell.value = int(num_val) if num_val.is_integer() else num_val
                    if col_idx <= len(headers) and ("valor" in headers[col_idx-1].lower() or "total" in headers[col_idx-1].lower() or "r$" in headers[col_idx-1].lower()):
                        cell.number_format = 'R$ #,##0.00'
                else:
                    cell.alignment = Alignment(horizontal="left", vertical="center")

        # Auto-fit Column Widths
        for col in ws.columns:
            max_len = max(len(str(cell.value or '')) for cell in col)
            col_letter = get_column_letter(col[0].column)
            ws.column_dimensions[col_letter].width = max(max_len + 4, 14)

        wb.save(filepath)
        return filepath

    # =========================================================================
    # 3. STRATEGY C: PDF (ReportLab com Capa, Sumário e Numeração de Páginas)
    # =========================================================================
    def compile_pdf(self, markdown_text: str, filepath: str, title: str = "Relatório Técnico BORAX") -> str:
        """
        Compiles Markdown into an elegant PDF document using ReportLab.
        - Cover Page (Capa) with Title, Subtitle, Author, Date and BORAX Theme.
        - Flowable Content with custom ParagraphStyles.
        - Footer with NumberedCanvas ('Página X de Y').
        """
        if not REPORTLAB_AVAILABLE:
            raise ImportError("reportlab não está instalado no ambiente backend Python.")

        doc = SimpleDocTemplate(
            filepath,
            pagesize=A4,
            leftMargin=54,  # ~1.9 cm
            rightMargin=54,
            topMargin=54,
            bottomMargin=54
        )

        styles = getSampleStyleSheet()

        # Custom Styles
        title_style = ParagraphStyle(
            'CoverTitle',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=24,
            leading=30,
            textColor=colors.HexColor('#14111F'),
            spaceAfter=15,
            alignment=0
        )

        subtitle_style = ParagraphStyle(
            'CoverSubtitle',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=13,
            leading=17,
            textColor=colors.HexColor('#6366F1'),
            spaceAfter=30,
            alignment=0
        )

        h1_style = ParagraphStyle(
            'CustomH1',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=16,
            leading=20,
            textColor=colors.HexColor('#14111F'),
            spaceBefore=18,
            spaceAfter=10,
            keepWithNext=True
        )

        h2_style = ParagraphStyle(
            'CustomH2',
            parent=styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=12,
            leading=15,
            textColor=colors.HexColor('#3F3F46'),
            spaceBefore=14,
            spaceAfter=6,
            keepWithNext=True
        )

        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontName='Helvetica',
            fontSize=10,
            leading=14,
            textColor=colors.HexColor('#27272A'),
            spaceAfter=8,
            alignment=4  # Justified
        )

        bullet_style = ParagraphStyle(
            'CustomBullet',
            parent=body_style,
            leftIndent=15,
            firstLineIndent=-10,
            spaceAfter=4
        )

        story = []

        # --- COVER PAGE ---
        story.append(Spacer(1, 80))
        story.append(Paragraph("PLATAFORMA DE IA LOCAL BORAX", ParagraphStyle('SubHeader', fontName='Helvetica-Bold', fontSize=10, textColor=colors.HexColor('#6366F1'), spaceAfter=15)))
        story.append(Paragraph(title.upper(), title_style))
        story.append(Paragraph("Relatório Técnico e Análise Aprofundada", subtitle_style))
        story.append(HRFlowable(width="100%", thickness=3, color=colors.HexColor('#6366F1'), spaceBefore=10, spaceAfter=30))

        story.append(Spacer(1, 140))
        story.append(Paragraph("<b>Autor:</b> Engine Local de IA BORAX", ParagraphStyle('Meta', fontName='Helvetica', fontSize=10, textColor=colors.HexColor('#52525B'), spaceAfter=4)))
        story.append(Paragraph("<b>Formato:</b> Documento Técnico Compilado (PDF)", ParagraphStyle('Meta2', fontName='Helvetica', fontSize=10, textColor=colors.HexColor('#52525B'), spaceAfter=4)))
        story.append(Paragraph("<b>Classificação:</b> Uso Interno / Acadêmico", ParagraphStyle('Meta3', fontName='Helvetica', fontSize=10, textColor=colors.HexColor('#52525B'))))
        
        story.append(PageBreak())

        # --- MAIN CONTENT FLOW ---
        lines = markdown_text.split("\n")
        in_table = False
        table_rows = []

        def flush_pdf_table(buffer):
            if not buffer:
                return
            t_data = []
            for line in buffer:
                if line.strip().startswith("|") and "---" not in line:
                    cells = [c.strip() for c in line.strip().split("|")[1:-1]]
                    if cells:
                        t_data.append([Paragraph(c, body_style) for c in cells])
            if t_data:
                t = Table(t_data, colWidths=None)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#14111F')),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ('TOPPADDING', (0, 0), (-1, -1), 6),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#E4E4E7')),
                    ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8F8FA')])
                ]))
                story.append(Spacer(1, 8))
                story.append(t)
                story.append(Spacer(1, 10))

        for line in lines:
            stripped = line.strip()

            if stripped.startswith("|"):
                in_table = True
                table_rows.append(stripped)
                continue
            else:
                if in_table:
                    flush_pdf_table(table_rows)
                    table_rows = []
                    in_table = False

            if not stripped:
                continue

            if stripped.startswith("# "):
                story.append(Paragraph(stripped[2:].strip(), h1_style))
                story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#E4E4E7'), spaceBefore=2, spaceAfter=8))
            elif stripped.startswith("## "):
                story.append(Paragraph(stripped[3:].strip(), h2_style))
            elif stripped.startswith("### "):
                story.append(Paragraph(stripped[4:].strip(), ParagraphStyle('H3', parent=h2_style, fontSize=11, fontName='Helvetica-BoldOblique')))
            elif stripped.startswith("- ") or stripped.startswith("* "):
                clean_text = stripped[2:].strip()
                clean_text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', clean_text)
                story.append(Paragraph(f"• {clean_text}", bullet_style))
            else:
                clean_text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', stripped)
                story.append(Paragraph(clean_text, body_style))

        if in_table:
            flush_pdf_table(table_rows)

        doc.build(story, canvasmaker=NumberedCanvas)
        return filepath
