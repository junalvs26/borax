import os
import shutil
import tempfile
import uuid
from datetime import datetime
from typing import Dict, Any, Optional, List

BORAX_EXPORTS_DIR = os.path.expanduser("~/.borax/exports")

def ensure_exports_dir():
    os.makedirs(BORAX_EXPORTS_DIR, exist_ok=True)

def generate_downloadable_file(content: str, file_format: str = "docx", filename: Optional[str] = None) -> Dict[str, Any]:
    """
    Compile text content into a downloadable document (.docx, .txt, .csv, .md)
    and save it in ~/.borax/exports/.
    """
    ensure_exports_dir()
    fmt = file_format.lower().replace(".", "")
    clean_name = filename or f"documento_borax_{str(uuid.uuid4())[:8]}"
    if not clean_name.endswith(f".{fmt}"):
        out_filename = f"{clean_name}.{fmt}"
    else:
        out_filename = clean_name

    out_path = os.path.join(BORAX_EXPORTS_DIR, out_filename)

    try:
        if fmt == "docx":
            import docx
            import re
            doc = docx.Document()
            
            lines = content.split("\n")
            for line in lines:
                stripped = line.strip()
                if not stripped:
                    continue
                if stripped.startswith("# "):
                    doc.add_heading(stripped[2:].strip(), level=1)
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:].strip(), level=2)
                elif stripped.startswith("### "):
                    doc.add_heading(stripped[4:].strip(), level=3)
                elif stripped.startswith("- ") or stripped.startswith("* "):
                    p = doc.add_paragraph(style='List Bullet')
                    text = stripped[2:].strip()
                    # Parse bold runs inside list item
                    parts = re.split(r'(\*\*.*?\*\*)', text)
                    for part in parts:
                        if part.startswith("**") and part.endswith("**"):
                            p.add_run(part[2:-2]).bold = True
                        else:
                            p.add_run(part)
                else:
                    p = doc.add_paragraph()
                    # Parse bold runs inside paragraph
                    parts = re.split(r'(\*\*.*?\*\*)', stripped)
                    for part in parts:
                        if part.startswith("**") and part.endswith("**"):
                            p.add_run(part[2:-2]).bold = True
                        else:
                            p.add_run(part)
            doc.save(out_path)
        elif fmt in ["txt", "md", "csv"]:
            with open(out_path, "w", encoding="utf-8") as f:
                f.write(content)
        else:
            # Fallback text file
            out_filename = f"{clean_name}.txt"
            out_path = os.path.join(BORAX_EXPORTS_DIR, out_filename)
            with open(out_path, "w", encoding="utf-8") as f:
                f.write(content)

        file_size = os.path.getsize(out_path)
        return {
            "status": "success",
            "filename": out_filename,
            "file_format": fmt,
            "filepath": out_path,
            "size_bytes": file_size,
            "download_url": f"/api/file/download/{out_filename}",
            "message": f"Arquivo '{out_filename}' gerado com sucesso!"
        }
    except Exception as e:
        print(f"[ChatPowerExecutor Error] Falha ao gerar arquivo: {e}")
        return {
            "status": "error",
            "message": f"Erro ao gerar arquivo {fmt}: {str(e)}"
        }

from intent_router import detect_intent, IntentType

class ChatPowerExecutor:
    def __init__(self, rag_engine=None, knowledge_manager=None, media_processor=None, data_analyst=None, ollama_service=None):
        self.rag_engine = rag_engine
        self.knowledge_manager = knowledge_manager
        self.media_processor = media_processor
        self.data_analyst = data_analyst
        self.ollama_service = ollama_service

    def execute_inline(
        self,
        power_trigger: str = "",
        text_input: str = "",
        file_path: Optional[str] = None,
        file_name: Optional[str] = None,
        file_format: Optional[str] = "docx"
    ) -> Dict[str, Any]:
        """
        Execute a power inline inside chat stream using 3-layer Intent Detection.
        """
        intent_info = detect_intent(
            message=text_input,
            file_name=file_name,
            active_power_override=power_trigger
        )
        intent = intent_info["intent"]

        if intent == IntentType.CHAT_ONLY:
            return {
                "status": "success",
                "power": "CHAT_ONLY",
                "type": "chat_only",
                "intent_info": intent_info,
                "message": "Mensagem classificada como diálogo comum (CHAT_ONLY)."
            }

        elif intent == IntentType.MOUNT_CARTRIDGE:
            if not file_path:
                return {"status": "error", "message": "Nenhum arquivo fornecido para criar a base de conhecimento."}
            res = self.knowledge_manager.mount_media(file_path, file_name=file_name)
            return {
                "status": "success",
                "power": "@criar-base",
                "type": "cartridge_mounted",
                "message": f"⚡ Base de Conhecimento **'{res['cartridge']['name']}'** criada e encaixada no leitor com sucesso!",
                "data": res
            }

        elif intent == IntentType.INGEST_FILE:
            if not file_path:
                return {"status": "error", "message": "Nenhum arquivo fornecido para ingestão."}
            res = self.rag_engine.process_file(file_path)
            return {
                "status": "success",
                "power": "@ingestao",
                "type": "file_ingested",
                "message": f"⚡ Arquivo **'{file_name or os.path.basename(file_path)}'** ingerido no LanceDB ({res.get('chunks_count', 0)} chunks vetoriais).",
                "data": res
            }

        elif intent == IntentType.TRANSCRIBE_MEDIA:
            if file_path:
                res = self.media_processor.transcribe_media(file_path=file_path)
            elif text_input and ("http://" in text_input or "https://" in text_input):
                res = self.media_processor.transcribe_media(url=text_input.strip())
            else:
                return {"status": "error", "message": "Envie um arquivo de áudio/vídeo ou URL do YouTube para transcrição."}
            
            transcript = res.get("full_transcript", "")
            return {
                "status": "success",
                "power": "@transcrever",
                "type": "transcription",
                "message": f"⚡ Mídia transcrita via Whisper ({res.get('duration', 0):.1f}s):\n\n```text\n{transcript[:1500]}\n```",
                "data": res
            }

        elif intent == IntentType.ANALYZE_DATA:
            if not file_path:
                return {"status": "error", "message": "Envie uma planilha ou arquivo CSV/Parquet para análise."}
            schema_res = self.data_analyst.inspect_schema(file_path)
            return {
                "status": "success",
                "power": "@analisar-dados",
                "type": "data_analysis",
                "message": f"⚡ Esquema da planilha **'{file_name or os.path.basename(file_path)}'** analisado via DuckDB:\n- Colunas: `{list(schema_res.get('columns', {}).keys())}`\n- Amostra processada com sucesso.",
                "data": schema_res
            }

        elif intent == IntentType.GENERATE_DOCUMENT:
            from document_wizard_engine import DocumentWizardEngine
            wiz = DocumentWizardEngine(ollama_service=self.ollama_service, rag_engine=self.rag_engine)
            res = wiz.generate_custom_document(
                prompt=text_input or "Estudo Técnico",
                preferences={
                    "theme": text_input or "Estudo Técnico",
                    "document_type": "abnt",
                    "tone": "scientific",
                    "export_format": file_format or "docx"
                }
            )
            return {
                "status": "success",
                "power": "@gerar-documento",
                "type": "downloadable_file",
                "message": res.get("message") or f"📄 Documento **'{res['file']['filename']}'** gerado com sucesso!",
                "document_text": res.get("document_text"),
                "file": res.get("file")
            }

        else:
            return {
                "status": "success",
                "power": "CHAT_ONLY",
                "type": "chat_only",
                "message": "Diálogo comum ativado."
            }
