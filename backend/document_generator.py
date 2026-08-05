import os
import re
import docx
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def convert_markdown_to_docx(markdown_text: str, filepath: str) -> str:
    """
    Convert LLM-synthesized Markdown text into an ABNT formatted Word (.docx) document.
    ABNT Rules:
    - Margins: Superior 3cm, Esquerda 3cm, Inferior 2cm, Direita 2cm.
    - Font: Arial 12pt.
    - Paragraph Alignment: Justified.
    - Line Spacing: 1.5.
    - First Line Indent: 1.25cm.
    """
    doc = docx.Document()

    # Set ABNT Page Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(3.0)
        section.left_margin = Cm(3.0)
        section.bottom_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Set Normal Paragraph Style (ABNT Body Text)
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    
    style_normal.paragraph_format.line_spacing = 1.5
    style_normal.paragraph_format.space_after = Pt(6)
    style_normal.paragraph_format.first_line_indent = Cm(1.25)
    style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Process Markdown lines
    lines = markdown_text.split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("# "):
            h1_text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(h1_text.upper())
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)

        elif stripped.startswith("## "):
            h2_text = stripped[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(h2_text)
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)

        elif stripped.startswith("### "):
            h3_text = stripped[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(h3_text)
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.italic = True
            run.font.color.rgb = RGBColor(0, 0, 0)

        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            text = stripped[2:].strip()
            
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.font.name = 'Arial'
                    r.font.size = Pt(12)
                    r.bold = True
                else:
                    r = p.add_run(part)
                    r.font.name = 'Arial'
                    r.font.size = Pt(12)

        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.font.name = 'Arial'
                    r.font.size = Pt(12)
                    r.bold = True
                else:
                    r = p.add_run(part)
                    r.font.name = 'Arial'
                    r.font.size = Pt(12)

    doc.save(filepath)
    return filepath
