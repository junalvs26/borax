import os
import re
import uuid
from typing import List, Dict, Any, Optional
from pathlib import Path
from pypdf import PdfReader
import docx
import lancedb
from sentence_transformers import SentenceTransformer

# Store LanceDB in ~/.app_data/lancedb
HOME_DIR = Path.home()
DB_DIR = HOME_DIR / ".app_data" / "lancedb"
DEFAULT_TABLE = "knowledge_base"

class RAGEngine:
    def __init__(self, db_dir: Optional[Path] = None):
        self.db_dir = db_dir or DB_DIR
        os.makedirs(self.db_dir, exist_ok=True)
        self.db = lancedb.connect(str(self.db_dir))
        # Embedding model local e leve
        self.embedder = SentenceTransformer("all-MiniLM-L6-v2")

    def _get_table_names(self) -> List[str]:
        """Safely fetch table names list from LanceDB connection."""
        try:
            res = self.db.table_names()
            if isinstance(res, (list, tuple)):
                return [str(x) for x in res]
            elif hasattr(res, "tables"):
                return [t.name if hasattr(t, "name") else str(t) for t in res.tables]
            return list(res)
        except Exception:
            return []

    def _ensure_table(self, table_name: str = DEFAULT_TABLE):
        """Ensures the specified LanceDB table exists."""
        tables = self._get_table_names()
        if table_name not in tables:
            sample_emb = self.embedder.encode("initial initialization text").tolist()
            dummy_data = [{
                "id": "init",
                "text": "Initialization vector space",
                "filename": "system",
                "vector": sample_emb
            }]
            return self.db.create_table(table_name, data=dummy_data)
        else:
            return self.db.open_table(table_name)

    def extract_text(self, file_path: str) -> str:
        """Extract text from .pdf, .txt, or .docx file."""
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        elif ext == ".pdf":
            reader = PdfReader(file_path)
            pages_text = []
            for page in reader.pages:
                txt = page.extract_text()
                if txt:
                    pages_text.append(txt)
            return "\n".join(pages_text)
        elif ext == ".docx":
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n".join(paragraphs)
        else:
            raise ValueError(f"Formato de arquivo não suportado: {ext}. Suportados: .pdf, .txt, .docx")

    def chunk_text(self, text: str, min_chars: int = 500, max_chars: int = 1000, overlap: int = 100) -> List[str]:
        """Split text into chunks of 500-1000 characters with 100 character overlap."""
        if not text or not text.strip():
            return []

        chunks = []
        start = 0
        text_len = len(text)

        while start < text_len:
            end = start + max_chars
            if end >= text_len:
                chunk = text[start:]
                if chunk.strip():
                    chunks.append(chunk.strip())
                break

            # Try to break chunk at a natural boundary (newline, period, space)
            break_pos = text.rfind('\n', start + min_chars, end)
            if break_pos == -1:
                break_pos = text.rfind('. ', start + min_chars, end)
            if break_pos == -1:
                break_pos = text.rfind(' ', start + min_chars, end)
            if break_pos == -1 or break_pos <= start:
                break_pos = end

            chunk = text[start:break_pos].strip()
            if chunk:
                chunks.append(chunk)

            # Move start pointer back by overlap
            start = max(start + 1, break_pos - overlap)

        return chunks

    def process_file(self, file_path: str, table_name: str = DEFAULT_TABLE) -> Dict[str, Any]:
        """Extract, chunk, embed, and store file content into LanceDB."""
        filename = os.path.basename(file_path)
        full_text = self.extract_text(file_path)
        chunks = self.chunk_text(full_text)

        if not chunks:
            return {"status": "error", "message": f"Nenhum texto extraído de '{filename}'."}

        embeddings = self.embedder.encode(chunks)
        table = self._ensure_table(table_name)

        records = []
        for chunk, emb in zip(chunks, embeddings):
            records.append({
                "id": str(uuid.uuid4()),
                "text": chunk,
                "filename": filename,
                "vector": emb.tolist()
            })

        table.add(records)
        return {
            "status": "success",
            "filename": filename,
            "table_name": table_name,
            "chunks_count": len(records),
            "message": f"Arquivo '{filename}' processado ({len(records)} chunks adicionados à tabela '{table_name}')."
        }

    def query_context(
        self,
        query: str,
        table_name: Any = DEFAULT_TABLE,
        top_k: int = 3,
        min_score: float = 0.45
    ) -> List[Dict[str, Any]]:
        """
        Hybrid Search & Re-ranking (Vector Similarity + Text Keyword Boost).
        Calculates similarity score and re-ranks context chunks above min_score.
        """
        tables_to_query = []
        if isinstance(table_name, list):
            tables_to_query = table_name
        elif isinstance(table_name, str):
            tables_to_query = [table_name]

        existing_tables = self._get_table_names()
        valid_tables = [t for t in tables_to_query if t in existing_tables]

        if not valid_tables:
            return []

        query_clean = (query or "").lower().strip()
        query_words = set(re.findall(r'\b\w{4,}\b', query_clean))

        query_emb = self.embedder.encode(query).tolist()
        combined_results = []

        for tbl_name in valid_tables:
            try:
                table = self.db.open_table(tbl_name)
                res = table.search(query_emb).limit(top_k * 2).to_list()
                for r in res:
                    if r.get("id") != "init":
                        r["source_table"] = tbl_name
                        
                        dist = r.get("_distance", 1.0)
                        # Normalize L2 distance to similarity score (0.0 to 1.0)
                        vector_score = max(0.0, 1.0 - (float(dist) / 2.0))
                        
                        # Hybrid Keyword Boost
                        text_lower = (r.get("text") or "").lower()
                        keyword_matches = sum(1 for w in query_words if w in text_lower)
                        keyword_boost = min(0.3, keyword_matches * 0.08)
                        
                        combined_score = vector_score + keyword_boost
                        r["similarity_score"] = round(combined_score, 4)
                        
                        if combined_score >= min_score:
                            combined_results.append(r)
            except Exception as e:
                print(f"[RAGEngine Error] Falha ao consultar tabela '{tbl_name}': {e}")

        # Re-rank by combined similarity score in descending order
        combined_results.sort(key=lambda x: x.get("similarity_score", 0.0), reverse=True)

        return combined_results[:top_k]

    def get_all_records(self, table_name: str = DEFAULT_TABLE) -> List[Dict[str, Any]]:
        """Fetch all records from a table."""
        tables = self._get_table_names()
        if table_name not in tables:
            return []

        table = self.db.open_table(table_name)
        records = table.search().limit(10000).to_list()
        return [r for r in records if r.get("id") != "init"]
